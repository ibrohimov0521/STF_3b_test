import argparse
import csv
import json
import re
from pathlib import Path

from database import AnswerOption, Question, get_session, init_db


def read_json(path: Path) -> list[dict]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, list):
        raise ValueError("JSON file must contain a list of tests")
    return data


def read_csv(path: Path) -> list[dict]:
    rows = []
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            rows.append(
                {
                    "question": row.get("question") or row.get("savol") or row.get("text"),
                    "info": row.get("info") or row.get("izoh") or row.get("explanation"),
                    "answers": [
                        row.get("answer1") or row.get("a") or row.get("javob1"),
                        row.get("answer2") or row.get("b") or row.get("javob2"),
                        row.get("answer3") or row.get("c") or row.get("javob3"),
                        row.get("answer4") or row.get("d") or row.get("javob4"),
                    ],
                }
            )
    return rows


def read_txt(path: Path) -> list[dict]:
    content = path.read_text(encoding="utf-8-sig").strip()
    return parse_text_blocks(content)


def read_docx(path: Path) -> list[dict]:
    from docx import Document

    document = Document(path)
    table_tests = read_docx_tables(document)
    if table_tests:
        return table_tests

    content = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return parse_text_blocks(content)


def read_docx_tables(document) -> list[dict]:
    tests = []
    current = None
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            cells += [""] * (7 - len(cells))
            number, _, question, answer1, answer2, answer3, answer4 = cells[:7]

            is_header = "Test topshirig" in question or "javob" in answer1.lower()
            has_number = bool(re.search(r"\d+", number))
            answers = [answer1, answer2, answer3, answer4]

            if is_header:
                continue

            if has_number and question and all(answers):
                current = {"question": question, "answers": answers}
                tests.append(current)
                continue

            if current is None:
                continue

            if question and not has_number:
                current["question"] = f"{current['question']} {question}".strip()

            for index, extra in enumerate(answers):
                if extra:
                    current["answers"][index] = f"{current['answers'][index]} {extra}".strip()

    return tests


def parse_text_blocks(content: str) -> list[dict]:
    blocks = re.split(r"\n\s*\n", content)
    tests = []
    for block in blocks:
        lines = [clean_line(line) for line in block.splitlines() if clean_line(line)]
        if len(lines) < 5:
            continue
        info = "\n".join(lines[5:]).strip() or None
        tests.append({"question": lines[0], "answers": lines[1:5], "info": info})
    return tests


def clean_line(line: str) -> str:
    line = line.strip()
    line = re.sub(r"^\d+[\).\-\s]+", "", line)
    line = re.sub(r"^[A-Da-d][\).\-\s]+", "", line)
    return line.strip()


def normalize_tests(raw_tests: list[dict]) -> list[dict]:
    normalized = []
    for index, item in enumerate(raw_tests, start=1):
        question = (item.get("question") or item.get("savol") or item.get("text") or "").strip()
        info = (item.get("info") or item.get("izoh") or item.get("explanation") or None)
        info = str(info).strip() if info else None
        answers = item.get("answers") or item.get("javoblar") or item.get("options") or []
        answers = [str(answer).strip() for answer in answers if answer and str(answer).strip()]
        if not question or len(answers) < 4:
            raise ValueError(f"Test #{index} has no question or fewer than 4 answers")
        normalized.append({"question": question, "answers": answers[:4], "info": info})
    return normalized


def load_tests(path: Path) -> list[dict]:
    suffix = path.suffix.lower()
    if suffix == ".json":
        raw = read_json(path)
    elif suffix == ".csv":
        raw = read_csv(path)
    elif suffix == ".txt":
        raw = read_txt(path)
    elif suffix == ".docx":
        raw = read_docx(path)
    else:
        raise ValueError("Supported formats: .json, .csv, .txt, .docx")
    return normalize_tests(raw)


def import_tests(path: Path, replace: bool) -> int:
    init_db()
    tests = load_tests(path)
    with get_session() as session:
        if replace:
            session.query(AnswerOption).delete()
            session.query(Question).delete()

        for item in tests:
            question = Question(text=item["question"], info=item.get("info"), source=path.name)
            question.options = [
                AnswerOption(text=answer, is_correct=(index == 0))
                for index, answer in enumerate(item["answers"])
            ]
            session.add(question)
        session.commit()
    return len(tests)


def main() -> None:
    parser = argparse.ArgumentParser(description="Import tests into the bot database.")
    parser.add_argument("file", help="Path to .json, .csv, .txt, or .docx test file")
    parser.add_argument("--replace", action="store_true", help="Delete old tests before importing")
    args = parser.parse_args()

    count = import_tests(Path(args.file), args.replace)
    print(f"Imported {count} tests")


if __name__ == "__main__":
    main()
